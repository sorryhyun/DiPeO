"""DOCX file format handler."""

from typing import Any
from pathlib import Path

import aiofiles


try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocxHandler:
    """Handler for DOCX files."""
    
    @property
    def supported_extensions(self) -> list[str]:
        """List of supported file extensions."""
        return [".docx", ".DOCX"] if HAS_DOCX else []
    
    def can_handle(self, file_path: Path) -> bool:
        """Check if this handler can process the given file.
        
        Args:
            file_path: Path to check
            
        Returns:
            True if this is a DOCX file and python-docx is available
        """
        return HAS_DOCX and file_path.suffix.lower() == ".docx"
    
    def read(self, file_path: Path) -> tuple[str, str]:
        """Read DOCX file and return (parsed_content, raw_content).
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (text content, text content)
            
        Raises:
            ImportError: If python-docx is not installed
        """
        if not HAS_DOCX:
            raise ImportError("python-docx is not installed")
        
        doc = Document(str(file_path))
        content = "\n".join([para.text for para in doc.paragraphs])
        # For DOCX, parsed and raw content are the same
        return content, content
    
    async def write(self, file_path: Path, content: Any) -> None:
        """Write content to DOCX file.
        
        Args:
            file_path: Path to write to
            content: Content to write (will be converted to text)
            
        Raises:
            ImportError: If python-docx is not installed
        """
        if not HAS_DOCX:
            raise ImportError("python-docx is not installed")
        
        doc = Document()
        
        # Convert content to string if needed
        text = str(content) if not isinstance(content, str) else content
        
        # Split by newlines and add paragraphs
        for line in text.split("\n"):
            doc.add_paragraph(line)
        
        # Save synchronously (python-docx doesn't support async)
        doc.save(str(file_path))
    
    def format_content(self, content: Any) -> str:
        """Format content as string.
        
        Args:
            content: Content to format
            
        Returns:
            String representation of content
        """
        return str(content) if not isinstance(content, str) else content